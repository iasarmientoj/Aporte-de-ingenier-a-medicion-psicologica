import pandas as pd
from docx import Document
from docx.shared import Pt
import os

def generar_examenes():
    nombre_archivo = "Banco de items.xlsx"
    
    # Verificar si el archivo existe
    if not os.path.exists(nombre_archivo):
        print(f"Error: No se encuentra el archivo '{nombre_archivo}' en la carpeta actual.")
        return

    print("Leyendo base de datos...")
    
    try:
        # Cargar las hojas del Excel
        # Hoja 'items': Debe tener columnas N, PREGUNTA, OPC A, OPC B, OPC C, OPC D
        df_items = pd.read_excel(nombre_archivo, sheet_name='items')
        
        # Hoja 'orden': Debe tener columnas N, TIPO A, TIPO B, TIPO C, TIPO D
        df_orden = pd.read_excel(nombre_archivo, sheet_name='orden')
        
        # Establecemos la columna 'N' como índice en items para buscar rápido por ID
        df_items.set_index('N', inplace=True)
        
    except Exception as e:
        print(f"Error al leer el Excel: {e}")
        return

    # Lista de los tipos de prueba (nombres de las columnas en la hoja 'orden')
    tipos_prueba = ['TIPO A', 'TIPO B', 'TIPO C', 'TIPO D']

    print("Generando documentos...")

    for tipo in tipos_prueba:
        # Crear un nuevo documento de Word
        doc = Document()
        
        # --- Encabezado del Examen ---
        titulo = doc.add_heading(f'PRUEBA DE CONOCIMIENTOS - {tipo}', 0)
        titulo.alignment = 1 # Centrado
        
        doc.add_paragraph('Nombre: _______________________________________________________')
        doc.add_paragraph('Fecha:  _________________________\n')
        
        intro = doc.add_paragraph('Instrucciones: Por favor responde las siguientes 20 preguntas de cultura general. Selecciona la opción correcta llenando completamente el círculo correspondiente de la hoja de respuestas.')
        intro.alignment = 1 # Centrado
        intro.style = 'Intense Quote'
        

        # --- Generación de Preguntas ---
        # Iteramos sobre la hoja 'orden' (que tiene los números del 1 al 20)
        for index, row in df_orden.iterrows():
            
            numero_visual = row['N']  # El número que verá el alumno (1, 2, 3...)
            id_pregunta = row[tipo]   # El ID real de la pregunta en el banco (sacado de la columna TIPO X)
            
            # Buscamos la info de esa pregunta específica en el banco de ítems
            try:
                item = df_items.loc[id_pregunta]
                texto_pregunta = item['PREGUNTA']
                opciones = {
                    'a': item['OPC A'],
                    'b': item['OPC B'],
                    'c': item['OPC C'],
                    'd': item['OPC D']
                }
            except KeyError:
                print(f"Advertencia: El ID {id_pregunta} no existe en la hoja de items.")
                continue

            # Escribir la pregunta en el Word
            # Usamos negrita para el enunciado
            p = doc.add_paragraph()
            run = p.add_run(f"{numero_visual}. {texto_pregunta}")
            run.bold = True
            run.font.size = Pt(11)

            # Escribir las opciones
            for letra, texto_opcion in opciones.items():
                p_opc = doc.add_paragraph(f"    {letra}) {texto_opcion}")
                p_opc.paragraph_format.space_after = Pt(2) # Reducir espacio entre opciones

            doc.add_paragraph('') # Espacio extra entre preguntas

        # --- Guardar el archivo ---
        nombre_salida = f"Prueba_Cultura_General_{tipo.replace(' ', '_')}.docx"
        doc.save(nombre_salida)
        print(f"Generado exitosamente: {nombre_salida}")

    print("\n¡Proceso terminado! Se han creado los 4 archivos Word.")

if __name__ == "__main__":
    generar_examenes()